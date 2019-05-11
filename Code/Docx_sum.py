from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import re
import pandas as pd
import glob
import os.path
from time import gmtime, strftime
import sys
import Kb_conn as kc
import time
import TextSumtry as ts
import csv
import os

# Create a function to find document title
def find_title(paragraphs):
    for paragraph in paragraphs:
        # Extract style Heading 1 as title
        if paragraph.style.name == 'Heading 1':
            return paragraph

#
# Create a function to identify the sequence ID of the headings
def get_headerID(para):
    # Create a list
    paraID = []
    # Get paragraph numbers
    paraNo = len(para)
    # Looping through the texts and find Heading 2, 3 and 4 by the style
    for i in range(paraNo):
        paragraph = para[i]
        if paragraph.style.name in ['Heading 2', 'Heading 3', 'Heading 4']:
            # Add the sequence ID of the headers to the list
            paraID.append(i)
    return paraID

# Create a function to include further info link
def add_pagelink(answer, page, link = "further"):
    if answer.strip() != '':
        if link == "further":   # For normal tailing link
            if page not in answer:
                # Processing answer string
                answer = answer.strip() + "\n\n\nFor " + link + " information: [" + page + "](" + page.replace(' ', '%20') + ")"
        else:   # link is "table"
            answer = answer.strip() + "\n\n\nFor " + link + " information: [" + page + "](" + page.replace(' ', '%20') + ")"
    return answer

# Create a function to extract Q&A
def getQA (para, headerID, title, ques, ans, exc, textsum, page):
    # Create global answer
    pg_answer = ""
    # Create global summarisation notice flag
    pg_sim_flag = False
    # Loop through all the paragraphs
    for i in range(len(para)):
        # Find headers by comparing sequence id with the headerID list, and exclude the one having the same text as title
        if i in headerID:
            # Create an answer string
            answer = ""
            # Create summarisation notice flag
            sim_flag = False
            # Create question variable
            h = para[i].text
            # Set numbers for potential Numbered List
            count = 1
            # Check if the header is similar to the page title or header is introduction
            if h.lower() == title.replace('&', 'and').lower() or any(s in h.lower() for s in exc):
                # Loop through the instances after the identified header
                for j in range(i + 1, len(para)):
                    # When getting another header, break to process next header
                    if j in headerID:
                        break
                    # When the instance is a paragraph
                    if isinstance(para[j], Paragraph):
                        # When getting style of List Paragraph, append the texts with bullet points
                        if para[j].style.name == 'List Paragraph':
                            # Set pre-text notation
                            pre = "* "
                            #print(para[j].text + " " + str(j) + " " + para[j].style.name + " " + para[i].text)
                            # Check and append the answer
                            pg_answer, pg_sim_flag = check_ts(para[j].text, textsum, pg_answer, pg_sim_flag, pre)
                            # Prevent messing up the Numbered List
                            count = 1
                        # When getting style of Numbered List, append the texts with numbers
                        elif para[j].style.name == 'Numbered List':
                            # Set pre-text notation
                            pre = str(count) + ". "
                            #print(para[i].text + " " + para[j].text)
                            # Check and append the answer
                            pg_answer, pg_sim_flag = check_ts(para[j].text, textsum, pg_answer, pg_sim_flag, pre)
                            # Count incremented by 1
                            count += 1
                        # Append the texts of other styles of paragraphs
                        else:
                            #print(para[j].text + " " + str(j) + " " +para[j].style.name +" " +para[i].text)
                            # Check and append the answer
                            pg_answer, pg_sim_flag = check_ts(para[j].text, textsum, pg_answer, pg_sim_flag)
                            # Prevent messing up the Numbered List
                            count = 1
                    # When the instance is a table
                    if isinstance(para[j], Table):
                        # Create a Not Null list
                        notnullinx = []
                        #print(len(para[j].columns))
                        # Loop through each column
                        for i in range(len(para[j].columns)):
                            #print(i)
                            #print(len(para[j].columns[i].cells))
                            # Checking how many columns are fully filled with contents
                            if not any(colcell.text.strip() == '' for colcell in para[j].columns[i].cells):
                                #print("not null!!")
                                # Record the index of full columns
                                notnullinx.append(i)
                                #print(notnullinx)
                        # When only 1 column has full contents
                        if len(notnullinx) == 1:
                            # Loop through the column's cells
                            for cell in para[j].columns[notnullinx[0]].cells:
                                # Loop through the paragraphs in the cell
                                for cellpara in cell.paragraphs:
                                    # When getting style of List Paragraph, append the texts with bullet points
                                    if cellpara.style.name == 'List Paragraph':
                                        # Set pre-text notation
                                        pre = "* "
                                        # print(para[j].text + " " + str(j) + " " + para[j].style.name + " " + para[i].text)
                                        # Check and append the answer
                                        pg_answer, pg_sim_flag = check_ts(cellpara.text, textsum, pg_answer, pg_sim_flag,
                                                                          pre)
                                        # Prevent messing up the Numbered List
                                        count = 1
                                    # When getting style of Numbered List, append the texts with numbers
                                    elif cellpara.style.name == 'Numbered List':
                                        # Set pre-text notation
                                        pre = str(count) + ". "
                                        # print(para[i].text + " " + para[j].text)
                                        # Check and append the answer
                                        pg_answer, pg_sim_flag = check_ts(cellpara.text, textsum, pg_answer, pg_sim_flag,
                                                                          pre)
                                        # Count incremented by 1
                                        count += 1
                                    # Append the texts of other styles of paragraphs
                                    else:
                                        # print(para[j].text + " " + str(j) + " " +para[j].style.name +" " +para[i].text)
                                        # Check and append the answer
                                        pg_answer, pg_sim_flag = check_ts(cellpara.text, textsum, pg_answer, pg_sim_flag)
                                        # Prevent messing up the Numbered List
                                        count = 1
                        # When table having more than 1 full columns
                        if len(notnullinx) > 1:
                            # Add table link for table having more than 2 columns
                            pganswer = add_pagelink(pganswer, page, "detailed table")
            else:
                #print(para[i].text.lower())
                #print(title.replace('&', 'and').lower())
                # Loop through the instances after the identified header
                for j in range(i+1, len(para)):
                    # When getting another header, break to process next header
                    if j in headerID:
                        break
                    # When the instance is a paragraph
                    if isinstance(para[j], Paragraph):
                        # When getting style of List Paragraph, append the texts with bullet points
                        if para[j].style.name == 'List Paragraph':
                            # Set pre-text notation
                            pre = "* "
                            # Check and append the answer
                            answer, sim_flag = check_ts(para[j].text, textsum, answer, sim_flag, pre)
                            # Prevent messing up the Numbered List
                            count = 1
                        # When getting style of Numbered List, append the texts with numbers
                        elif para[j].style.name == 'Numbered List':
                            # Set pre-text notation
                            pre = str(count) + ". "
                            # Check and append the answer
                            answer, sim_flag = check_ts(para[j].text, textsum, answer, sim_flag, pre)
                            # Count incremented by 1
                            count += 1
                        # Append the texts of other styles of paragraphs
                        else:
                            # Check and append the answer
                            answer, sim_flag = check_ts(para[j].text, textsum, answer, sim_flag)
                            # Prevent messing up the Numbered List
                            count = 1
                    if isinstance(para[j], Table):
                        # Create a Not Null list
                        notnullinx = []
                        #print(len(para[j].columns))
                        # Loop through each column
                        for i in range(len(para[j].columns)):
                            #print(i)
                            #print(len(para[j].columns[i].cells))
                            # Checking how many columns are fully filled with contents
                            if not any(colcell.text.strip() == '' for colcell in para[j].columns[i].cells):
                                #print("not null!!")
                                # Record the index of full columns
                                notnullinx.append(i)
                                #print(notnullinx)
                        # When only 1 column has full contents
                        if len(notnullinx) == 1:
                            # Loop through the column's cells
                            for cell in para[j].columns[notnullinx[0]].cells:
                                # Loop through the paragraphs in the cell
                                for cellpara in cell.paragraphs:
                                    # When getting style of List Paragraph, append the texts with bullet points
                                    if cellpara.style.name == 'List Paragraph':
                                        # Set pre-text notation
                                        pre = "* "
                                        # print(para[j].text + " " + str(j) + " " + para[j].style.name + " " + para[i].text)
                                        # Check and append the answer
                                        answer, sim_flag = check_ts(cellpara.text, textsum, answer, sim_flag,
                                                                          pre)
                                        # Prevent messing up the Numbered List
                                        count = 1
                                    # When getting style of Numbered List, append the texts with numbers
                                    elif cellpara.style.name == 'Numbered List':
                                        # Set pre-text notation
                                        pre = str(count) + ". "
                                        # print(para[i].text + " " + para[j].text)
                                        # Check and append the answer
                                        answer, sim_flag = check_ts(cellpara.text, textsum, answer, sim_flag,
                                                                          pre)
                                        # Count incremented by 1
                                        count += 1
                                    # Append the texts of other styles of paragraphs
                                    else:
                                        # print(para[j].text + " " + str(j) + " " +para[j].style.name +" " +para[i].text)
                                        # Check and append the answer
                                        answer, sim_flag = check_ts(cellpara.text, textsum, answer, sim_flag)
                                        # Prevent messing up the Numbered List
                                        count = 1
                        # When table having more than 1 full columns
                        if len(notnullinx) > 1:
                            # Add table link for table having more than 2 columns
                            answer = add_pagelink(answer, page, "detailed table")
                        #print(para[j].columns[2].text)
                        #colno = len(para[j].columns)
                        #rowno = len(para[j].rows)
                        #
                        # for i, row in enumerate(para[j].columns):
                        #     print(i)
                answer = style_repro(answer)
                answer = add_pagelink(answer, page)
                # Concat title and question and add to question list
                ques.append(title + " " + h.strip())
                # Add corresponding answer to answer list
                ans.append(answer.strip())
    return ques, ans, pg_answer

# Create a function looping through to get all files
def get_files(files):
    # Retrieve all docx files
    for file_name in glob.glob("*.docx"):
        # Put the file names into a list
        files.append(file_name)
        #print(file_name)


# Create a function to replicate the links
def style_repro(text):
    #print(text)
    # Find links in the Answer
    sub = re.findall("https?:\/\/?[\w/\-?=%.]+\.[\w/\-?=%]+", text)
    # If links found
    if len(sub) != 0:
        #print(sub)
        # Loop through found links
        for url in sub:
            # Replace the links with adding formatting syntax
            #print(url)
            text = re.sub(url, "[" + url + "](" + url + ")", text)
        #print(text)

    # Find mails in the Answer
    sub = re.findall(r"(^[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+$)", text)
    # If links found
    if len(sub) != 0:
        #print(sub)
        # Loop through found mails
        for mail in sub:
            # Replace the mails with adding formatting syntax
            text = re.sub(mail, "[" + mail + "](" + mail + ")", text)
        #print(text)

    return text

def iter_block_items(parent):
    """
    Generate a reference to each paragraph and table child within *parent*,
    in document order. Each returned value is an instance of either Table or
    Paragraph. *parent* would most commonly be a reference to a main
    Document object, but also works for a _Cell object, which itself can
    contain paragraphs and tables.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
        # print(parent_elm.xml)
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

# Create a function to check the text if needs summarisation
def check_ts(text, textsum, answer, sim_flag, pre = ""):
    # When the textsum is assigned TRUE
    if textsum:
        # When the text has a greater wordcount than 80
        if len(text.split()) > 80:
            # Call the TextSum and summarise the text
            text = ts.text_sum(text)
            # Set the summarisation notice to TRUE and add notice to the beginning of the answer
            if sim_flag == False:
                answer = "The following content has been summarised...\n\n\n" + answer
                sim_flag = True
        # Collate the answer texts
    answer = answer + pre + text + "\n\n"
    # Return answer and flag for other paragraphs
    return answer, sim_flag

# Create a function to output data frame to Q&A list - OLD
# def output_df(ques, ans, path):
#     # Create a data frame with the Q and A lists
#     df = pd.DataFrame({'Question': ques, 'Answer': ans, 'Source': "", 'Metadata': ""})
#     # Output the data frame to tsv file
#     df.to_csv(path + "\kk.tsv", sep='\t', encoding='utf-8', index=False, quoting=3, escapechar='\\')
#     # Return file path for post-processing
#     return path + "\kk.tsv"

# Create a function to make the file into QnA maker accepted formats - OLD
# def post_proce(file, newfile):
#     # Open the outputted file
#     with open(file, "rt", encoding='utf-8') as fin:
#         # Create another file
#         with open(newfile, "wt", encoding='utf-8') as fout:
#             # Parse line-by-line with replacing all the \n into raw \n\n
#             for line in fin:
#                 fout.write(line.replace('\\\n', r'\n\n'))
#     # Remove originally outputted file
#     os.remove(file)

# Create docx main function
def docx_main(ques, ans, dir_path, textsum):
    # Set directory path storing the DOCX files
    #dir_path = r"C:\Users\JamesC\OneDrive - Queensland University of Technology\IFN701Project\Deliverable\Code\docx" #sys.argv[2]   # SSIS var
    # Set output directory path
    #dest_path = r"C:\Users\JamesC\OneDrive - Queensland University of Technology\IFN701Project\Deliverable"  #sys.argv[1] # SSIS var
    # Set document directory
    os.chdir(dir_path)
    # Create a file, question and answer list
    files = []
    # Read from thr csv file into list for the links
    data = dict(csv.reader(open('C:\\Users\\jchou\\Desktop\\Docxlinks.csv')))

    # Create a list of excluding keywords
    excwords = ['further', 'useful link', 'intro']

    #question = []
    #answer = []

    # Get all files
    get_files(files)

    # Loop through each file to extract Q&A
    for file in files:
        # Get docx link from csv
        page = data[file]
        # Get the document for the library
        doc = Document(dir_path+"\\"+file)
        if (len(doc.paragraphs)>1):
            # Get document title
            title = find_title(doc.paragraphs).text.capitalize()
            # Get paragraphs and tables into list
            # para = list(doc.paragraphs)
            para = list(iter_block_items(doc))
            # Get header IDs in paragraph list
            headerID = get_headerID(para)
            # Extract Q&A
            ques, ans, pg_answer = getQA(para, headerID, title, ques, ans, excwords, textsum, page)
            # Add global page Q&A
            if pg_answer != "":
                pg_answer = style_repro(pg_answer)
                pg_answer = add_pagelink(pg_answer, page)
                # Add title to question list
                ques.append(title)
                # Add corresponding answer to answer list
                ans.append(pg_answer.strip())
        print(dir_path+"\\"+file + " processed......" )
    return ques, ans, len(files)




### Without API nee belows

# Output Q&A to tsv file
#file = output_df(question, answer, dest_path)
#jsonbody = output_df(question, answer, dest_path)
# Set another path for post-processing
#new = dest_path + "\\" + strftime("%Y%m%d%H%M%S", gmtime()) + "docx.tsv"
# Execute post-processing
#post_proce(file, new)


#print(question)
#print(answer[-1])
#print(len(question))
#print(len(answer))


# Update knowledge base with QnAMaker API
#kc.kb_update(jsonbody)
# Publish knowledge base with QnAMaker API
#kc.kb_publish()