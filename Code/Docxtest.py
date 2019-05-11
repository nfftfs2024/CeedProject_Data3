from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

import pandas as pd
import glob
import os.path
from time import gmtime, strftime
import sys
import Kb_conn as kc
import time
import csv
import os

# Create a function to include further info link
def add_pagelink(answer, page, link = "further"):
    if answer.strip() != '':
        if link == "further":   # For normal tailing link
            if page not in answer:
                # Processing answer string
                answer = answer.strip() + "\n\n\nFor " + link + " information: " + page.replace(' ', '%20')
        else:   # link is "table"
            answer = answer.strip() + "\n\n\nFor " + link + " information: " + page.replace(' ', '%20')
    return answer
# Create a function to extract Q&A
def getQA (para, headerID, title, ques, ans, exc, page):
    # Create global answer
    pg_answer = ""
    # Loop through all the paragraphs
    for i in range(len(para)):
        # Find headers by comparing sequence id with the headerID list, and exclude the one having the same text as title
        if i in headerID:
            # Create an answer string
            answer = ""
            # Create question variable
            h = para[i].text
            # Set numbers for potential Numbered List
            count = 1
            # Check if the header is similar to the page title or header is introduction
            if h.lower() == title.replace('&', 'and').lower() or any(s in h.lower() for s in exc):
                for j in range(i + 1, len(para)):
                    # When getting another header, break to process next header
                    if j in headerID:
                        break
                    if isinstance(para[j], Paragraph):
                        # When getting style of List Paragraph, append the texts with bullet points
                        if para[j].style.name == 'List Paragraph':
                            #print(para[j].text + " " + str(j) + " " + para[j].style.name + " " + para[i].text)
                            pg_answer = pg_answer + "* " + para[j].text + "\n\n"
                            # Prevent messing up the Numbered List
                            count = 1
                        # When getting style of Numbered List, append the texts with numbers
                        elif para[j].style.name == 'Numbered List':
                            #print(para[i].text + " " + para[j].text)
                            pg_answer = pg_answer + str(count) + ". " + para[j].text + "\n\n"
                            # Count incremented by 1
                            count += 1
                        # Append the texts of other styles of paragraphs
                        else:
                            #print(para[j].text + " " + str(j) + " " +para[j].style.name +" " +para[i].text)
                            pg_answer = pg_answer + para[j].text + "\n\n"
                            # Prevent messing up the Numbered List
                            count = 1
                    if isinstance(para[j], Table):
                        print("table la")
            else:
                #print(para[i].text.lower())
                #print(title.replace('&', 'and').lower())
                # Loop through the paragraphs after the identified header
                for j in range(i+1, len(para)):
                    # When getting another header, break to process next header
                    if j in headerID:
                        break
                    if isinstance(para[j], Paragraph):
                        # When getting style of List Paragraph, append the texts with bullet points
                        if para[j].style.name == 'List Paragraph':
                            answer = answer + "* " + para[j].text + "\n\n"
                            # Prevent messing up the Numbered List
                            count = 1
                        # When getting style of Numbered List, append the texts with numbers
                        elif para[j].style.name == 'Numbered List':
                            answer = answer + str(count) + ". " + para[j].text + "\n\n"
                            # Count incremented by 1
                            count += 1
                        # Append the texts of other styles of paragraphs
                        else:
                            answer = answer + para[j].text + "\n\n"
                            # Prevent messing up the Numbered List
                            count = 1
                    # If the block is table
                    if isinstance(para[j], Table):
                        notnullinx = []
                        #print(len(para[j].columns))
                        for i in range(len(para[j].columns)):
                            #print(i)
                            #print(len(para[j].columns[i].cells))
                            if not any(s.text.strip() == '' for s in para[j].columns[i].cells):
                                #print("not null!!")
                                notnullinx.append(i)
                                print(notnullinx)
                        # Based on the columns numbers
                        if len(notnullinx) == 1:
                            for x in para[j].columns[notnullinx[0]].cells:
                                for y in x.paragraphs:
                                    if y.style.name == 'List Paragraph':
                                        # print(para[j].text + " " + str(j) + " " + para[j].style.name + " " + para[i].text)
                                        answer = answer + "* " + y.text + "\n\n"
                                        # Prevent messing up the Numbered List
                                        count = 1
                                        # When getting style of Numbered List, append the texts with numbers
                                    elif y.style.name == 'Numbered List':
                                        # print(para[i].text + " " + para[j].text)
                                        answer = answer + str(count) + ". " + y.text + "\n\n"
                                        # Count incremented by 1
                                        count += 1
                                        # Append the texts of other styles of paragraphs
                                    else:
                                        # print(para[j].text + " " + str(j) + " " +para[j].style.name +" " +para[i].text)
                                        answer = answer + y.text + "\n\n"
                                        # Prevent messing up the Numbered List
                                        count = 1
                        #print(para[j].columns[2].text)
                        #colno = len(para[j].columns)
                        #rowno = len(para[j].rows)
                        #
                        # for i, row in enumerate(para[j].columns):
                        #     print(i)
                answer = add_pagelink(answer, page)
                # Concat title and question and add to question list
                ques.append(title + " " + h.strip())
                # Add corresponding answer to answer list
                ans.append(answer.strip())
    return ques, ans, pg_answer

# Create a function looping through to get all files
def get_files(files):
    # Retrieve all docx files
    for file_name in glob.glob("*.docx"):
        # Put the file names into a list
        files.append(file_name)
        #print(file_name)

# Create a function to find document title
def find_title(paragraphs):
    for paragraph in paragraphs:
        # Extract style Heading 1 as title
        if paragraph.style.name == 'Heading 1':
            return paragraph

# Create a function to identify the sequence ID of the headings
def get_headerID(para):
    # Create a list
    paraID = []
    # Get paragraph numbers
    paraNo = len(para)
    # Looping through the texts and find Heading 2, 3 and 4 by the style
    for i in range(paraNo):
        paragraph = para[i]
        if paragraph.style.name in ['Heading 2', 'Heading 3', 'Heading 4']:
            # Add the sequence ID of the headers to the list
            #print(i)
            #print(paragraph.text)
            paraID.append(i)
    return paraID

def iter_block_items(parent):
    """
    Generate a reference to each paragraph and table child within *parent*,
    in document order. Each returned value is an instance of either Table or
    Paragraph. *parent* would most commonly be a reference to a main
    Document object, but also works for a _Cell object, which itself can
    contain paragraphs and tables.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
        # print(parent_elm.xml)
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

# Create docx main function
def docx_main(ques, ans, dir_path):
    # Set directory path storing the DOCX files
    #dir_path = r"C:\Users\JamesC\OneDrive - Queensland University of Technology\IFN701Project\Deliverable\Code\docx" #sys.argv[2]   # SSIS var
    # Set output directory path
    #dest_path = r"C:\Users\JamesC\OneDrive - Queensland University of Technology\IFN701Project\Deliverable"  #sys.argv[1] # SSIS var
    # Set document directory
    os.chdir(dir_path)
    # Create a file, question and answer list
    files = []
    # Read from thr csv file into list for the links
    data = dict(csv.reader(open('D:\\vsyo.csv')))

    # Create a list of excluding keywords
    excwords = ['further', 'useful link', 'intro']

    #question = []
    #answer = []

    # Get all files
    get_files(files)

    # Loop through each file to extract Q&A
    for file in files:
        # Get docx link from csv
        page = data[file]
        # Get the document for the library
        doc = Document(dir_path+"\\"+file)
        # Get document title
        title = find_title(doc.paragraphs).text.capitalize()
        # para = list(doc.paragraphs)
        para = list(iter_block_items(doc))
        # Get header IDs in paragraph list
        headerID = get_headerID(para)
        # Extract Q&A
        ques, ans, pg_answer = getQA(para, headerID, title, ques, ans, excwords, page)

        # Add global page Q&A
        if pg_answer != "":
            pg_answer = add_pagelink(pg_answer, page)
            # Add title to question list
            ques.append(title)
            # Add corresponding answer to answer list
            ans.append(pg_answer.strip())
        print(dir_path + "\\" + file + " processed......")
    return ques, ans




    #     print(block.style.name)
    # elif isinstance(block, Table):
    #     for i, row in enumerate(block.rows):
    #         text = (cell.text for cell in row. cells)
    #         if i == 0:
    #             keys = tuple(text)
    #             continue
    #
    #         row_data = dict(zip(keys, text))
    #         print(str(row_data))

# table = doc.tables
# print(table)
# para = list(doc.paragraphs)
# print(len(para))
# for t in table:
#     for c in t.rows:
#         print(c.cells[2].text)
# # Put paragraphs into list
# para = list(doc.paragraphs)
# # Get header IDs in paragraph list
# headerID = get_headerID(para)
# # Extract Q&A
# ques, ans, pg_answer = getQA(para, headerID, title, ques, ans, excwords, page)
# # Add global page Q&A
# if pg_answer != "":
#     pg_answer = add_pagelink(pg_answer, page)
#     # Add title to question list
#     ques.append(title)
#     # Add corresponding answer to answer list
#     ans.append(pg_answer.strip())
# print(dir_path+"\\"+file + " processed......" )
# for table in doc.tables:
#     for row in table.rows:
#         for cell in row.cells:
#             print(cell.text)



